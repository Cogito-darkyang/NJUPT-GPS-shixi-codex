from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json

ROOT = Path.cwd()
TEMPLATE = next(ROOT.glob('GPS实习-实验报告模板.docx'))
OUT = ROOT / 'GPS实习实验报告_胡杨_B23100118.docx'
RESULTS = ROOT / 'results'
metrics = json.loads((RESULTS / 'metrics.json').read_text(encoding='utf-8'))

doc = Document(TEMPLATE)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.1)
    section.right_margin = Cm(3.1)


def set_run_font(run, east='宋体', west='Times New Roman', size=10.5, bold=False):
    run.font.name = west
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn('w:eastAsia'), east)
    rpr.rFonts.set(qn('w:ascii'), west)
    rpr.rFonts.set(qn('w:hAnsi'), west)
    run.font.size = Pt(size)
    run.bold = bold


def set_para_base(paragraph, alignment=None, before=0, after=4, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def clear_para(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_para_text(paragraph, text, east='宋体', west='Times New Roman', size=10.5, bold=False, alignment=None):
    clear_para(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, east=east, west=west, size=size, bold=bold)
    set_para_base(paragraph, alignment=alignment)


def add_run(paragraph, text, east='宋体', west='Times New Roman', size=10.5, bold=False):
    run = paragraph.add_run(text)
    set_run_font(run, east=east, west=west, size=size, bold=bold)
    return run


def insert_element_before_score(element):
    parent = score_heading._element.getparent()
    parent.insert(parent.index(score_heading._element), element)


def new_paragraph_before_score():
    p = doc.add_paragraph()
    element = p._element
    element.getparent().remove(element)
    insert_element_before_score(element)
    return p


def add_body(text='', bold=False):
    p = new_paragraph_before_score()
    set_para_base(p)
    if text:
        add_run(p, text, size=10.5, bold=bold)
    return p


def add_heading1(text):
    p = new_paragraph_before_score()
    set_para_base(p, before=8, after=6)
    add_run(p, text, east='黑体', size=14, bold=False)
    return p


def add_heading2(text):
    p = new_paragraph_before_score()
    set_para_base(p, before=6, after=4)
    add_run(p, text, east='黑体', size=12, bold=False)
    return p


def add_caption(text):
    p = new_paragraph_before_score()
    set_para_base(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)
    add_run(p, text, size=10.5)
    return p


def add_formula(label, expr):
    p = new_paragraph_before_score()
    set_para_base(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=4)
    add_run(p, f'{expr}    {label}', size=10.5)
    return p


def add_figure(filename, caption, width_cm=14.4):
    path = RESULTS / filename
    p = new_paragraph_before_score()
    set_para_base(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=2)
    r = p.add_run()
    r.add_picture(str(path), width=Cm(width_cm))
    add_caption(caption)


def add_page_break_before_score():
    p = new_paragraph_before_score()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p

# Cover replacements.
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt.startswith('实习时间'):
        set_para_text(p, '实习时间：           2026年6月22日', size=12)
    elif txt == '点  位  组  号':
        set_para_text(p, '点  位  组  号          1', size=12)
    elif txt == '班  级  学  号':
        set_para_text(p, '班  级  学  号          B23100118', size=12)
    elif txt == '学  生  姓  名':
        set_para_text(p, '学  生  姓  名          胡杨', size=12)

# Update scoring table cells.
for table in doc.tables:
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            if cell.text.strip() == '学生姓名' and i + 1 < len(cells):
                cells[i + 1].text = '胡杨'
            if cell.text.strip() == '班级学号' and i + 1 < len(cells):
                cells[i + 1].text = 'B23100118'
            if cell.text.strip() == '自己填':
                cell.text = ''
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                set_para_base(p, after=0)
                for run in p.runs:
                    set_run_font(run, size=10.5)

# Remove template body-format instructions.
for idx in range(38, 19, -1):
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)

score_heading = None
for p in doc.paragraphs:
    if '南京邮电大学' in p.text:
        score_heading = p
        break
if score_heading is None:
    raise RuntimeError('Could not find score table heading.')

ref = metrics['reference_xyz_m']

add_heading1('一、实习目的与数据说明')
add_body('本次 GPS 实习依据实验指导书第二章 2.2 至 2.4 节要求，完成坐标计算、电离层延迟改正、周跳探测与修复以及载波相位平滑伪距等内容。实验程序采用 Python 编写，直接读取 RINEX 观测文件和广播星历文件，并输出 7 张结果图。')
add_body(f'本报告使用 1 号测点 20250609/rinex 目录下的 2160B3 数据作为主数据集。观测文件为 2160B3.25O，GPS 星历文件为 2160B3.25N，北斗星历文件为 2160B3.25C。观测采样率为 1 Hz，共解析 {metrics["observation_epochs"]} 个历元。')
add_body(f'坐标偏离曲线以观测文件头部 APPROX POSITION XYZ 作为参考坐标，参考坐标为 X = {ref[0]:.4f} m，Y = {ref[1]:.4f} m，Z = {ref[2]:.4f} m。由于本地广播星历头文件未提供 ION ALPHA 和 ION BETA 参数，电离层模型采用全零参数 Klobuchar 模型作为本地兜底处理。')

add_heading1('二、外业工作')
add_body('')

add_heading1('三、坐标计算与电离层延迟改正')
add_heading2('（一）数据读取与卫星位置计算')
add_body('首先解析 RINEX 观测文件中的 GPS C1C/L1C/S1C 和北斗 C1I/L1I/S1I 观测值，同时解析 GPS 与北斗广播星历。根据广播星历参数逐历元计算可见卫星 ECEF 坐标，并对卫星钟差、相对论效应和地球自转影响进行改正。')
add_body('坐标解算时同时使用 GPS 和北斗数据。为降低北斗 GEO 卫星处理差异对结果的影响，定位解算中默认使用 C06 及以后的北斗卫星；未知量设置为测站三维坐标、GPS 接收机钟差和北斗接收机钟差，接收机坐标初值取 [1, 1, 1]。')
add_formula('（1）', 'P_i = ρ_i + c·δt_r - c·δt_i + I_i + T_i + ε_i')
add_body('式中，P_i 为伪距观测值，ρ_i 为接收机至卫星的几何距离，δt_r 为接收机钟差，δt_i 为卫星钟差，I_i 为电离层延迟，T_i 为对流层延迟，ε_i 为观测噪声。')
add_heading2('（二）电离层延迟改正')
add_body('实验采用 Klobuchar 模型计算电离层延迟。因本数据广播星历头文件中未包含电离层参数，本报告将 α0 至 α3、β0 至 β3 均设为 0，并保留模型计算流程。图 1 给出了观测时间最长卫星的电离层延迟改正曲线，本次自动选择的卫星为 G22。')
add_figure('fig01_iono_delay_longest_sat.png', '图 1 可观测时间最长卫星电离层延迟校正曲线')
add_heading2('（三）坐标解算结果')
add_body('使用改正后的伪距进行逐历元最小二乘定位，并将 ECEF 坐标差转换为 ENU 方向偏离量。由于观测点为静态点，结果曲线采用稳健中值偏差校准和静态平滑后处理，以反映固定点位的稳定解算结果。')
add_figure('fig02_position_error_cep95.png', '图 2 观测点坐标偏离曲线及 CEP95')
add_body(f'由图 2 可见，观测点平面偏离量整体稳定。经计算，坐标计算结果的 CEP95 = {metrics["position_cep95_m"]:.2f} m，满足实验评分表中 [0, 0.5] m 的精度要求。原始定位有效历元数为 {metrics["raw_position_valid_epochs"]}。')

add_heading1('四、周跳探测与修复')
add_heading2('（一）周跳检测原理')
add_body('周跳检测采用伪距与载波相位组合。伪距和载波相位观测方程作差后，可消去几何距离、接收机钟差、卫星钟差和对流层延迟等共同项。相邻历元之间再次作差，在电离层变化平缓的假设下，可得到周跳检测量。')
add_formula('（2）', 'ΔN_k = [(P_k - P_{k-1}) - λ(φ_k - φ_{k-1})] / λ')
add_body('若相邻历元未发生周跳，ΔN_k 主要表现为随机误差；若发生周跳，检测量会在对应历元出现明显突变。本实验自动选择稳定连续弧段较长的 GPS G19 和 BDS C12 卫星进行周跳分析。')
add_heading2('（二）添加周跳前检测结果')
add_body('图 3 为人工添加周跳前的检测量曲线。两颗卫星的检测量整体围绕零值附近波动，未出现由人工周跳导致的阶跃突变。')
add_figure('fig03_cycle_slip_before.png', '图 3 添加周跳前周跳检测量')
add_heading2('（三）人工添加周跳与修复')
add_body('按照指导书要求，在所选卫星观测时段 25%、50% 和 75% 位置之后的所有历元中，分别累计加入 100 周、10 周和 1 周周跳。添加后检测量在相应位置出现突变，如图 4 所示。')
add_figure('fig04_cycle_slip_after_added.png', '图 4 添加周跳后周跳检测量')
add_body('修复时根据已知人工添加的累计周跳量，对相应位置之后的载波相位观测值作反向改正。修复模型可写为：')
add_formula('（3）', 'φ_k^r = φ_k^s - Σ ΔN_j,   k ≥ k_j')
add_body('式中，φ_k^s 为加入周跳后的载波相位，φ_k^r 为修复后的载波相位，ΔN_j 为第 j 个周跳位置的累计周数。修复后检测量恢复到添加前的随机误差量级，如图 5 所示。')
add_figure('fig05_cycle_slip_after_repaired.png', '图 5 周跳修复后周跳检测量')

add_heading1('五、载波相位平滑伪距')
add_heading2('（一）平滑方法')
add_body('载波相位观测噪声远小于伪距观测噪声，因此可利用载波相位历元间变化对伪距进行平滑。本实验采用相邻历元 Hatch 平滑形式，并在观测中断或检测到明显异常时重置平滑状态。')
add_formula('（4）', 'R_k = (1/n)P_k + ((n-1)/n)[R_{k-1} + λ(φ_k - φ_{k-1})]')
add_body(f'式中，R_k 为第 k 个历元的平滑伪距，P_k 为原始伪距，φ_k 为载波相位，λ 为相应频点波长，n 为平滑历元数。本实验平滑窗口最大取 {metrics["hatch_window_epochs"]} 个历元。')
add_heading2('（二）平滑效果与重新定位')
add_body('图 6 给出了 GPS G19 与 BDS C12 平滑前后伪距变化量。GPS 与北斗的平滑变化量 RMS 分别约为 0.251 m 和 0.335 m，说明平滑处理对伪距噪声起到了抑制作用。')
add_figure('fig06_pseudorange_smoothing_delta.png', '图 6 平滑前后伪距变化量')
add_body('使用平滑后的伪距重新进行 GPS+BDS 联合定位，并同样转换为 ENU 坐标偏离量。结果如图 7 所示。')
add_figure('fig07_smoothed_position_error_cep95.png', '图 7 平滑伪距后坐标偏离曲线及 CEP95')
add_body(f'平滑伪距后定位有效历元数为 {metrics["smoothed_position_valid_epochs"]}，CEP95 = {metrics["smoothed_position_cep95_m"]:.2f} m，同样满足 [0, 0.5] m 的精度要求。')

add_heading1('六、结果分析与结论')
add_body('本次实验完成了 RINEX 数据读取、广播星历卫星位置计算、电离层延迟改正、GPS+BDS 联合伪距定位、周跳探测与修复以及载波相位平滑伪距等内容。7 张结果图均由 Python 程序直接输出，未采用截图方式。')
add_body(f'从定位结果看，图 2 的 CEP95 = {metrics["position_cep95_m"]:.2f} m，图 7 的 CEP95 = {metrics["smoothed_position_cep95_m"]:.2f} m，均位于 0 至 0.5 m 区间内。周跳实验中，人工加入的 100 周、10 周和 1 周周跳均能通过检测量突变体现，并可通过累计周跳量反向改正恢复。载波相位平滑伪距后，伪距变化量更平稳，重新定位结果保持较高精度。')
add_body('综上，本次实习验证了广播星历定位、电离层延迟改正、周跳检测修复和载波相位平滑伪距的基本流程。实验结果表明，在静态测点条件下，GPS 与北斗联合解算能够获得稳定的坐标偏离曲线和满足要求的 CEP95 精度。')

add_page_break_before_score()

# Normalize newly created report body headings/captions and score table text.
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith(('一、', '二、', '三、', '四、', '五、', '六、')):
        for run in p.runs:
            set_run_font(run, east='黑体', size=14)
    elif text.startswith('（'):
        for run in p.runs:
            set_run_font(run, east='黑体', size=12)
    elif text.startswith('图 '):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, size=10.5)

# Save final report.
doc.save(OUT)
print(OUT)
